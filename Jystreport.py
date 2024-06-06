import sys
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QLabel, QLineEdit, QTextEdit, QComboBox, 
    QPushButton, QVBoxLayout, QHBoxLayout, QWidget, QFileDialog, 
    QTableWidget, QTableWidgetItem, QMessageBox
)
from PyQt6.QtGui import QPixmap, QImage
from PyQt6.QtCore import Qt
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING

import os
import tempfile
from datetime import datetime

class VulnerabilityReportApp(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("简易版等保渗透报告生成器 by L1ech0 v1.5")
        self.setGeometry(100, 100, 100, 1200)

        self.vulnerabilities = []

        self.vulnerability_info = {
            "选择漏洞类型": {
                "name": "存在漏洞",
                "desc": "",
                "suggestion": ""
            },
            "接口泄露": {
                "name": "存在接口泄露漏洞",
                "desc": "1.未授权访问： 攻击者可以通过未受保护的接口访问敏感数据或功能，获取原本只有授权用户或系统才能访问的信息或操作权限。\n2.数据泄露： 泄漏的接口可能暴露用户的个人信息、业务数据、财务信息等敏感数据，导致隐私侵犯和经济损失。\n3.身份盗窃： 攻击者可以利用泄露的接口获取用户身份信息，用于冒充合法用户进行恶意活动，如欺诈、开设银行账户等。",
                "suggestion": "1.认证和授权： 确保所有接口都经过严格的认证和授权控制，只有合法用户才能访问和执行相应操作。\n2.最小权限原则： 实施最小权限原则，限制接口功能只对必需的权限开放，防止未授权的访问和操作。\n3.加密传输： 使用HTTPS、SSL/TLS等加密协议保护接口数据在传输过程中的安全，防止中间人攻击和数据窃取。4.\n输入验证和过滤： 对接口的输入数据进行严格验证和过滤，防止SQL注入、脚本注入等攻击。"
            },
            "SQL注入": {
                "name": "存在SQL注入漏洞",
                "desc": "1.未授权访问： 攻击者可以通过SQL注入访问本不应公开的数据库内容，例如用户信息、财务数据和商业机密等。/n2.数据泄露： 攻击者可以获取敏感数据，如用户名、密码、个人信息等，这可能导致隐私泄露和身份盗窃。\n3.数据篡改： 攻击者可以修改数据库中的数据，例如更改用户账户信息、交易记录、库存数据等，导致数据完整性受损。\n4.数据删除： 攻击者可以删除数据库中的数据，导致数据丢失，影响系统正常运行和业务操作。\n5.执行恶意操作： 攻击者可以利用SQL注入执行恶意操作，例如创建新用户账户、提升权限、执行系统命令等，进一步扩大攻击影响。",
                "suggestion": "1.使用参数化查询： 使用参数化查询（Prepared Statements）或存储过程代替直接拼接SQL语句，从根本上防止SQL注入。\n2.使用ORM框架： 使用对象关系映射（ORM）框架，通过ORM框架生成的SQL查询来避免手动拼接SQL语句，降低SQL注入风险。\n3.输入验证和过滤： 对用户输入的数据进行严格验证和过滤，确保输入的数据符合预期格式，避免恶意SQL代码注入。\n4.最小权限原则： 对数据库用户和应用程序用户实施最小权限原则，限制其只能执行必要的操作，减少潜在的攻击面。\n5.错误信息处理： 不要在前端显示详细的数据库错误信息，避免攻击者通过错误信息获取数据库结构和敏感信息。"
            },
            "未授权访问": {
                "name": "存在未授权访问漏洞",
                "desc": "1.数据泄露： 攻击者可以访问和窃取敏感数据，包括个人信息、财务数据、商业机密等，导致隐私侵犯和经济损失。\n2.数据篡改： 未授权用户可能篡改或删除数据，破坏数据的完整性和准确性，影响业务操作和决策。\n3.滥用系统资源： 攻击者可以利用未授权访问执行高权限操作，消耗系统资源，导致合法用户无法正常使用系统。",
                "suggestion": "1.强制访问控制： 实施严格的访问控制策略，确保只有授权用户才能访问受保护的资源和数据。\n2.身份验证： 使用多因素身份验证（MFA）增强用户身份验证的安全性，防止身份冒用。\n3.权限管理： 实施基于角色的访问控制（RBAC）或最小权限原则（PoLP），限制用户权限只授予其执行任务所需的最小权限。\n4.加密数据传输： 使用加密技术（如SSL/TLS）保护数据在传输过程中的安全，防止中间人攻击。"
            },
            "文件上传": {
                "name": "存在文件上传漏洞",
                "desc": "1.远程代码执行： 攻击者可以上传包含恶意代码的文件，并通过访问这些文件在服务器上执行任意代码，从而完全控制服务器。\n2.恶意文件传播： 攻击者可以上传并分发恶意软件、病毒、木马等，感染服务器和用户设备，进一步扩展攻击范围。\n3.服务器资源滥用： 攻击者可以上传大量大文件或恶意文件，消耗服务器存储和带宽资源，导致性能下降或拒绝服务（DoS）。\n4.信息泄露： 攻击者可以上传脚本文件，通过这些文件访问和窃取服务器上的敏感信息，如配置文件、数据库凭据等。\n5.网页篡改： 攻击者可以上传篡改后的网页文件，替换或修改原有网页内容，进行钓鱼攻击或传播虚假信息，损害企业声誉。",
                "suggestion": "1.文件类型验证： 严格验证上传文件的类型，仅允许特定类型的文件（如图片、文档等）上传，拒绝可执行文件和脚本文件。\n2.文件名和扩展名检查： 验证和规范化上传文件的文件名和扩展名，避免通过文件名欺骗绕过验证。\n3.文件内容扫描： 使用反病毒软件或内容扫描工具对上传文件进行扫描，检测并阻止包含恶意代码的文件。\n4.文件上传位置限制： 将上传文件存放在指定目录，并设置该目录为不可执行，防止上传的文件被直接访问和执行。\n5.随机重命名文件： 对上传的文件进行随机重命名，避免文件名冲突和路径遍历攻击。"
            },
            "命令执行": {
                "name": "存在命令执行漏洞",
                "desc": "1.远程代码执行： 攻击者可以通过漏洞在目标服务器上执行任意命令或代码，完全控制服务器。\n2.数据泄露： 攻击者可以执行命令访问和窃取服务器上的敏感数据，包括用户信息、数据库内容、配置文件等。\n3.数据篡改： 攻击者可以执行命令篡改或删除服务器上的数据，破坏数据的完整性和可用性。\n4.恶意软件传播： 攻击者可以通过命令执行漏洞上传和运行恶意软件，感染服务器并进一步扩展攻击。",
                "suggestion": "1.输入验证和过滤： 对用户输入进行严格验证和过滤，拒绝包含特殊字符或命令语法的输入。\n2.使用安全API： 避免直接使用系统命令执行函数（如system、exec等），而应使用更安全的API或库进行相关操作。\n3.最小权限原则： 运行应用程序和相关进程时采用最小权限原则，限制其执行系统命令的权限。\n4.参数化命令： 对于必须执行的系统命令，使用参数化的方法，将用户输入作为参数传递，避免直接拼接命令字符串。"
            },
            "验证码重复利用": {
                "name": "存在验证码重复利用漏洞",
                "desc": "1.绕过身份验证： 攻击者可以通过重复使用验证码绕过身份验证步骤，未经授权地访问受保护的资源或执行敏感操作。\n2.自动化攻击： 攻击者可以利用重复使用的验证码进行自动化攻击，如暴力破解攻击、大规模注册虚假账户、垃圾信息发布等，破坏系统的正常运行。\n3.账户劫持： 攻击者可以通过验证码绕过安全机制，劫持用户账户，获取账户中的敏感信息或进行恶意操作。",
                "suggestion": "1.验证码一次性使用： 确保每个验证码只能使用一次，验证成功后立即使其失效，防止重复利用。\n2.验证码过期时间： 设置验证码的有效期，超时后验证码自动失效，防止长时间内被重复使用。\n3.验证码唯一性： 生成足够随机且唯一的验证码，防止通过预测或重用验证码进行攻击。\n4.服务器端验证： 在服务器端验证验证码的有效性，避免在客户端进行验证，防止通过客户端绕过验证机制。\n5.操作绑定： 将验证码与特定操作或会话绑定，确保验证码只能用于预定操作，防止跨操作或跨会话使用。"
            },
            "任意文件下载": {
                "name": "存在任意文件下载漏洞",
                "desc": "1.敏感信息泄露： 攻击者可以下载服务器上的敏感文件，如配置文件、数据库备份、用户信息文件等，导致敏感数据泄露。\n2.系统配置曝光： 攻击者可以下载包含系统配置的文件，如/etc/passwd、/etc/shadow、web.config等，获取系统配置信息，进一步分析和利用漏洞。\n3.源代码泄露： 如果攻击者能够下载应用程序的源代码文件，将会了解应用的内部逻辑和实现，发现更多潜在漏洞，提高攻击成功率。\n4.凭证和密钥泄露： 攻击者可以下载包含凭证和密钥的文件，如数据库连接配置、API密钥等，使用这些信息进行进一步的攻击。",
                "suggestion": "1.输入验证和过滤： 对用户输入的文件路径和文件名进行严格验证和过滤，确保只能访问预定义的安全目录和文件。\n2.路径规范化： 对用户输入的文件路径进行规范化处理，防止通过路径遍历（如../）访问不应公开的文件。\n3.限制文件访问： 在服务器端设置文件访问权限，仅允许下载特定目录中的文件，防止访问系统敏感文件。\n4.白名单机制： 使用白名单机制，仅允许下载特定类型或特定目录下的文件，拒绝其他未经授权的文件访问请求。\n5.文件名和路径编码： 对用户输入的文件名和路径进行适当的编码和转义，防止特殊字符被解释为路径的一部分。"
            },
            "短信轰炸": {
                "name": "存在短信轰炸漏洞",
                "desc": "1.用户困扰： 大量的无用短信会极大地干扰用户的正常生活和工作，导致用户手机不断收到通知，影响用户体验。\n2.经济损失： 对于按短信数量收费的服务，短信轰炸会导致财务损失，用户可能需要支付高额的短信费用。\n3.资源消耗： 短信轰炸会大量消耗短信服务提供商的资源，如短信网关带宽、服务器处理能力，影响正常用户的服务质量。",
                "suggestion": "1.验证码限制： 对同一手机号码在一定时间内的短信发送次数进行限制，防止短时间内频繁发送短信。\n2.IP限制： 限制同一IP地址在短时间内的短信发送请求次数，防止通过单一IP地址进行短信轰炸。\n3.行为分析和监控： 通过分析用户行为识别异常的短信发送请求，及时发现并阻止短信轰炸攻击。"
            },
            "XSS（跨站脚本攻击）": {
                "name": "存在XSS（跨站脚本攻击）漏洞",
                "desc": "1.会话劫持：攻击者可以通过XSS窃取用户的会话Cookie，从而冒充用户进行操作，访问受保护的资源。\n2.身份盗用：攻击者可以伪装成受害用户，执行一些敏感操作，如修改账号信息、进行交易等，造成身份盗用。\n3.数据泄露：攻击者可以通过XSS攻击获取用户输入的敏感信息，如用户名、密码、信用卡信息等，导致数据泄露。\n4.恶意重定向：攻击者可以利用XSS漏洞将用户重定向到恶意网站，诱骗用户输入敏感信息或下载恶意软件。",
                "suggestion": "1.输入验证和过滤：对用户输入的数据进行严格验证和过滤，确保其不包含恶意脚本代码。可以使用白名单策略，只允许合法的输入字符和格式。\n2.输出编码：在将用户输入的内容输出到网页时，进行适当的编码（如HTML、JavaScript、CSS等），防止恶意脚本被执行。\n3.使用安全库和框架：使用成熟的安全库和框架，它们通常内置了防止XSS攻击的机制，如Django、Rails、ASP.NET等。"
            },
            "越权漏洞": {
                "name": "存在越权漏洞",
                "desc": "1.数据泄露：攻击者可以访问到不应被其查看的敏感数据，如用户个人信息、公司机密文件、财务数据等。\n2.数据篡改：攻击者能够修改、删除或添加数据，造成数据的不完整性和不准确性，可能导致严重的业务影响。\n3.账户劫持：攻击者可以访问或控制其他用户的账户，从而进行假冒或恶意操作，例如金融交易、身份盗用等。\n4.系统破坏：攻击者可以执行系统管理操作，如停止服务、修改配置、删除重要文件，导致系统崩溃或服务中断。",
                "suggestion": "1.角色和权限管理：实施严格的角色和权限管理，确保每个用户只能访问和操作其权限范围内的资源。\n2.最小权限原则：遵循最小权限原则（Principle of Least Privilege），只授予用户完成任务所需的最低权限。\n3.权限验证：在每个请求操作前进行权限验证，确保用户具有执行该操作的合法权限。\n4.访问控制列表（ACL）：使用访问控制列表（ACL）来定义和管理各资源的访问权限，确保权限分配的准确性和灵活性。"
            },
            "弱口令": {
                "name": "存在弱口令漏洞",
                "desc": "1.未授权访问： 攻击者可以通过暴力破解、字典攻击或其他猜测方法获取弱口令，从而获得未授权的访问权限。攻击者可以访问敏感数据、修改系统配置，甚至完全控制受影响的系统。\n2.数据泄露： 一旦攻击者获得了访问权限，他们可以访问和窃取敏感数据，例如个人信息、财务数据、知识产权等。这可能导致隐私泄露、财务损失和声誉损害。\n3.恶意操作： 攻击者可以执行恶意操作，如删除文件、修改数据、安装后门程序或恶意软件。这可能导致系统的不稳定，甚至被用作进一步攻击其他系统的跳板。",
                "suggestion": "1.强制密码策略： 实施强密码策略，要求密码包含大小写字母、数字和特殊字符，并且长度不少于规定的最小长度。\n2.定期更改密码： 要求用户定期更改密码，避免长期使用同一个密码。\n3.实施账户锁定： 在检测到多次失败的登录尝试后，暂时锁定账户，防止暴力破解。\n4.多因素认证（MFA）： 增加额外的认证步骤，如短信验证码、手机应用验证或生物识别，使得单一密码不足以完全验证用户身份。"
            }

        }

        self.initUI()

    def initUI(self):
        main_layout = QVBoxLayout()

        # 输入区域
        form_layout = QVBoxLayout()

        self.report_name = self.create_form_row(form_layout, "报告名称：")
        
        # 漏洞类型下拉框
        self.entry_type = self.create_form_row(form_layout, "漏洞类型：", QComboBox)
        self.entry_type.addItems(self.vulnerability_info.keys())
        self.entry_type.currentIndexChanged.connect(self.update_vulnerability_info)
        
        self.entry_name = self.create_form_row(form_layout, "漏洞名称：")
        self.entry_level = self.create_form_row(form_layout, "漏洞级别：", QComboBox)
        self.entry_level.addItems(["低危", "中危", "高危"])
        self.entry_addr = self.create_form_row(form_layout, "漏洞地址：")
        self.text_desc = self.create_form_row(form_layout, "漏洞描述：", QTextEdit)
        self.text_suggestion = self.create_form_row(form_layout, "修补建议：", QTextEdit)

        self.proof_entries = []
        self.add_proof_button(form_layout, "漏洞证明：", "粘贴剪切板图片", self.paste_clipboard_image)
        self.entry_figure = self.create_form_row(form_layout, "图 例：")
        self.add_figure_button(form_layout, "添加图例", self.add_figure)

        main_layout.addLayout(form_layout)

        # 按钮区域
        button_layout = QHBoxLayout()

        add_button = QPushButton("新增漏洞", self)
        add_button.clicked.connect(self.add_vulnerability)
        button_layout.addWidget(add_button)

        update_button = QPushButton("更新漏洞", self)
        update_button.clicked.connect(self.update_vulnerability)
        button_layout.addWidget(update_button)

        delete_button = QPushButton("删除漏洞", self)
        delete_button.clicked.connect(self.delete_vulnerability)
        button_layout.addWidget(delete_button)

        report_button = QPushButton("生成报告", self)
        report_button.clicked.connect(self.generate_report)
        button_layout.addWidget(report_button)

        main_layout.addLayout(button_layout)

        # 漏洞列表
        self.table = QTableWidget(self)
        self.table.setColumnCount(6)
        self.table.setHorizontalHeaderLabels(["漏洞名称", "漏洞级别", "漏洞地址", "漏洞描述", "修补建议", "图例"])
        self.table.itemSelectionChanged.connect(self.on_table_select)
        main_layout.addWidget(self.table)

        # 设置主窗口的中心组件
        container = QWidget()
        container.setLayout(main_layout)
        self.setCentralWidget(container)

    def create_form_row(self, layout, label_text, widget_cls=QLineEdit):
        layout_h = QHBoxLayout()
        label = QLabel(label_text, self)
        widget = widget_cls(self)
        layout_h.addWidget(label)
        layout_h.addWidget(widget)
        layout.addLayout(layout_h)
        return widget

    def add_proof_button(self, layout, label_text, button_text, button_callback):
        layout_h = QHBoxLayout()
        label = QLabel(label_text, self)
        self.proof_thumbnail = QLabel(self)  # 用于显示缩略图
        self.proof_thumbnail.setFixedSize(100, 100)
        self.proof_thumbnail.setStyleSheet("border: 1px solid black;")
        button = QPushButton(button_text, self)
        button.clicked.connect(button_callback)
        layout_h.addWidget(label)
        layout_h.addWidget(self.proof_thumbnail)
        layout_h.addWidget(button)
        layout.addLayout(layout_h)

    def add_figure_button(self, layout, button_text, button_callback):
        button = QPushButton(button_text, self)
        button.clicked.connect(button_callback)
        layout.addWidget(button)

    def add_vulnerability(self):
        name = self.entry_name.text()
        level = self.entry_level.currentText()
        addr = self.entry_addr.text()
        desc = self.text_desc.toPlainText().strip()
        suggestion = self.text_suggestion.toPlainText().strip()
        proof_entries = self.proof_entries
        v_type = self.entry_type.currentText()

        if not name or not level or not addr or not desc or not suggestion or not v_type:
            QMessageBox.warning(self, "警告", "所有字段都是必填项")
            return

        self.vulnerabilities.append({
            "漏洞名称": name,
            "漏洞级别": level,
            "漏洞地址": addr,
            "漏洞描述": desc,
            "修补建议": suggestion,
            "漏洞证明": proof_entries,
            "漏洞类型": v_type
        })

        self.clear_form()
        self.update_vulnerability_list()

    def update_vulnerability_list(self):
        self.table.setRowCount(len(self.vulnerabilities))
        for row, vuln in enumerate(self.vulnerabilities):
            self.table.setItem(row, 0, QTableWidgetItem(vuln["漏洞名称"]))
            self.table.setItem(row, 1, QTableWidgetItem(vuln["漏洞级别"]))
            self.table.setItem(row, 2, QTableWidgetItem(vuln["漏洞地址"]))
            self.table.setItem(row, 3, QTableWidgetItem(vuln["漏洞描述"]))
            self.table.setItem(row, 4, QTableWidgetItem(vuln["修补建议"]))
            self.table.setItem(row, 5, QTableWidgetItem(", ".join(entry["图例"] for entry in vuln["漏洞证明"])))

    def paste_clipboard_image(self):
        clipboard = QApplication.clipboard()
        mime_data = clipboard.mimeData()

        if mime_data.hasImage():
            image = clipboard.image()
            temp_path = tempfile.mktemp(suffix=".png")
            image.save(temp_path)
            self.proof_thumbnail.setPixmap(QPixmap.fromImage(image).scaled(100, 100, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation))
            self.proof_entries.append({"路径": temp_path, "图例": ""})
        else:
            QMessageBox.warning(self, "警告", "剪切板中没有图片")

    def add_figure(self):
        figure_text = self.entry_figure.text().strip()
        if figure_text and self.proof_entries:
            self.proof_entries[-1]["图例"] = figure_text
            self.entry_figure.clear()

    def on_table_select(self):
        selected_items = self.table.selectedItems()
        if not selected_items:
            return

        row = selected_items[0].row()
        vuln = self.vulnerabilities[row]

        self.entry_name.setText(vuln["漏洞名称"])
        self.entry_level.setCurrentText(vuln["漏洞级别"])
        self.entry_addr.setText(vuln["漏洞地址"])
        self.text_desc.setPlainText(vuln["漏洞描述"])
        self.text_suggestion.setPlainText(vuln["修补建议"])
        self.proof_entries = vuln["漏洞证明"]
        self.entry_type.setCurrentText(vuln["漏洞类型"])

        # 更新缩略图
        if self.proof_entries:
            pixmap = QPixmap(self.proof_entries[-1]["路径"]).scaled(100, 100, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
            self.proof_thumbnail.setPixmap(pixmap)
        else:
            self.proof_thumbnail.clear()

    def update_vulnerability(self):
        selected_items = self.table.selectedItems()
        if not selected_items:
            return

        row = selected_items[0].row()
        vuln = self.vulnerabilities[row]

        name = self.entry_name.text()
        level = self.entry_level.currentText()
        addr = self.entry_addr.text()
        desc = self.text_desc.toPlainText().strip()
        suggestion = self.text_suggestion.toPlainText().strip()
        proof_entries = self.proof_entries
        v_type = self.entry_type.currentText()

        if not name or not level or not addr or not desc or not suggestion or not v_type:
            QMessageBox.warning(self, "警告", "所有字段都是必填项")
            return

        vuln.update({
            "漏洞名称": name,
            "漏洞级别": level,
            "漏洞地址": addr,
            "漏洞描述": desc,
            "修补建议": suggestion,
            "漏洞证明": proof_entries,
            "漏洞类型": v_type
        })

        self.update_vulnerability_list()

    def delete_vulnerability(self):
        selected_items = self.table.selectedItems()
        if not selected_items:
            return

        row = selected_items[0].row()
        del self.vulnerabilities[row]
        self.update_vulnerability_list()
        self.clear_form()

    def set_paragraph_format(self, paragraph):
        # 设置段后间距为0磅
        paragraph.paragraph_format.space_after = Pt(0)
        # 设置多倍行距为1.25倍
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = 1.25

    def create_numbered_paragraph(self, document, text):
        p = document.add_paragraph(style='List Number')
        self.set_paragraph_format(p)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        return p

    def generate_report(self):
        if not self.vulnerabilities:
            QMessageBox.warning(self, "警告", "没有漏洞信息可生成报告")
            return

        document = Document()

        for index, vuln in enumerate(self.vulnerabilities, start=1):
            # 漏洞名称，三号加粗，宋体，前加编号
            p = self.create_numbered_paragraph(document, f'{vuln["漏洞名称"]}')

            # 漏洞级别，常量小四加粗，变量小四加粗红色字体
            p = document.add_paragraph()
            self.set_paragraph_format(p)
            run = p.add_run('漏洞级别：')
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run = p.add_run(vuln["漏洞级别"])
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run.font.color.rgb = RGBColor(255, 0, 0)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 漏洞地址，常量小四加粗，变量小四字体
            p = document.add_paragraph()
            self.set_paragraph_format(p)
            run = p.add_run('漏洞地址：')
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run = p.add_run(vuln["漏洞地址"])
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 漏洞描述，常量小四加粗，变量小四字体
            p = document.add_paragraph()
            self.set_paragraph_format(p)
            run = p.add_run('漏洞描述：')
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.add_run().add_break()
            run = p.add_run(f'\t{vuln["漏洞描述"]}')
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 修补建议，常量小四加粗，变量小四字体
            p = document.add_paragraph()
            self.set_paragraph_format(p)
            run = p.add_run('修补建议：')
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.add_run().add_break()
            run = p.add_run(f'\t{vuln["修补建议"]}')
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 漏洞证明，常量小四加粗
            p = document.add_paragraph()
            self.set_paragraph_format(p)
            run = p.add_run('漏洞证明：')
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            for proof in vuln["漏洞证明"]:
                if proof["路径"] and os.path.exists(proof["路径"]):
                    try:
                        document.add_picture(proof["路径"], width=Inches(4))
                    except Exception as e:
                        p = document.add_paragraph()
                        self.set_paragraph_format(p)
                        p.add_run(f'插入图片失败: {proof["路径"]}, 错误: {e}')
                
                # 图例，变量小四字体
                if proof["图例"]:
                    p = document.add_paragraph()
                    self.set_paragraph_format(p)
                    p.alignment = 1  # 居中对齐
                    run = p.add_run(proof["图例"])
                    run.font.size = Pt(12)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 添加一个空行
            document.add_paragraph('')

        report_name = self.report_name.text()
        if not report_name:
            report_name = "报告"

        # 添加今天日期到报告名称
        today = datetime.today().strftime('%Y-%m-%d')
        report_name_with_date = f"{report_name}_{today}"

        report_file, _ = QFileDialog.getSaveFileName(self, "保存报告文件", report_name_with_date, "Word Document (*.docx)")
        if report_file:
            try:
                document.save(report_file)
                QMessageBox.information(self, "成功", "报告生成成功")
            except Exception as e:
                QMessageBox.warning(self, "错误", f"报告生成失败: {e}")

    def clear_form(self):
        self.entry_name.clear()
        self.entry_level.setCurrentIndex(0)
        self.entry_addr.clear()
        self.text_desc.clear()
        self.text_suggestion.clear()
        self.proof_entries = []
        self.proof_thumbnail.clear()
        self.entry_figure.clear()
        self.entry_type.setCurrentIndex(0)

    def update_vulnerability_info(self):
        vul_type = self.entry_type.currentText()
        if vul_type in self.vulnerability_info:
            self.entry_name.setText(self.vulnerability_info[vul_type]["name"])
            self.text_desc.setPlainText(self.vulnerability_info[vul_type]["desc"])
            self.text_suggestion.setPlainText(self.vulnerability_info[vul_type]["suggestion"])

if __name__ == "__main__":
    app = QApplication(sys.argv)
    main_win = VulnerabilityReportApp()
    main_win.show()
    sys.exit(app.exec())
